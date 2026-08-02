import json
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_blocks(parent):
    if not isinstance(parent, DocumentType):
        raise TypeError("Only top-level document blocks are supported")

    for child in parent.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def paragraph_record(paragraph, index):
    properties = paragraph._p.pPr
    numbering = None
    if properties is not None and properties.numPr is not None:
        num_id = properties.numPr.numId
        level = properties.numPr.ilvl
        numbering = {
            "num_id": num_id.val if num_id is not None else None,
            "level": level.val if level is not None else None,
        }

    return {
        "kind": "paragraph",
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "numbering": numbering,
        "text": paragraph.text,
    }


def table_record(table, index):
    return {
        "kind": "table",
        "index": index,
        "rows": [
            [cell.text.replace("\r", "").strip() for cell in row.cells]
            for row in table.rows
        ],
    }


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: inspect_wy_docx.py INPUT.docx OUTPUT.json")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(input_path)
    blocks = []
    paragraph_index = 0
    table_index = 0

    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            blocks.append(paragraph_record(block, paragraph_index))
            paragraph_index += 1
        else:
            blocks.append(table_record(block, table_index))
            table_index += 1

    payload = {
        "source": str(input_path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "section_count": len(document.sections),
        "blocks": blocks,
    }
    output_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    headings = [
        block
        for block in blocks
        if block["kind"] == "paragraph"
        and block["text"].strip()
        and (
            "Heading" in (block["style"] or "")
            or "标题" in (block["style"] or "")
            or block["style"] in {"Title", "Subtitle"}
        )
    ]

    print(
        json.dumps(
            {
                "paragraph_count": payload["paragraph_count"],
                "table_count": payload["table_count"],
                "inline_shape_count": payload["inline_shape_count"],
                "section_count": payload["section_count"],
                "headings": headings,
                "tables": [
                    {
                        "index": block["index"],
                        "row_count": len(block["rows"]),
                        "column_count": max(
                            (len(row) for row in block["rows"]),
                            default=0,
                        ),
                        "first_row": block["rows"][0] if block["rows"] else [],
                    }
                    for block in blocks
                    if block["kind"] == "table"
                ],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
